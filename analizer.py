import string
import os
import re
import langid
import fitz
from collections import Counter
from ResultModel import ResultModel
from docx import Document
#from pyth.plugins.plaintext.writer import PlaintextWriter
#from pyth.plugins.rtf15.reader import Rtf15Reader

class TextAnalizatorCore:
    
    path_to_stopwords_file_ru = "stopwords_ru.txt"
    path_to_stopwords_file_ua = "stopwords_ua.txt"
    path_to_stopwords_file_eng = "stopwords_eng.txt"

    @staticmethod 
    def Analize(path,words,exclude):
        _, file_extension = os.path.splitext(path)
        if file_extension == ".txt":
            return TextAnalizatorCore.AnalizeTxt(path, words, exclude)
        elif file_extension == ".docx":
            return TextAnalizatorCore.AnalizeDocx(path, words, exclude)
        elif file_extension == ".rtf":
            return TextAnalizatorCore.AnalizeRtf(path, words, exclude)
        elif file_extension == ".pdf":
            return TextAnalizatorCore.AnalizePdf(path, words, exclude)
        elif file_extension == ".doc":
            return TextAnalizatorCore.AnalizeDoc(path, words, exclude)


    @staticmethod
    def AnalizePdf(path, words, exclude):
        try:
            text = ""
            with fitz.open(path) as pdf_document:
                num_pages = pdf_document.page_count
                for page_number in range(num_pages):
                    page = pdf_document[page_number]
                    text += page.get_text()
            return TextAnalizatorCore.AnalizeGeneric(text, words, exclude)
        except Exception as e:
            print(f"An error occurred: {e}")
            return None 

    @staticmethod
    def AnalizeTxt(path, words, exclude):
        try:
            with open(path, 'r',  encoding='utf-8') as file:
                text = file.read() 
                return TextAnalizatorCore.AnalizeGeneric(text, words, exclude)
        except Exception as e:
            print(f"An error occurred: {e}")
            return None   


    @staticmethod
    def AnalizeDoc(path, words, exclude):
        try:
            doc = Document(path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text = '\n'.join(full_text)
            return TextAnalizatorCore.AnalizeGeneric(text, words, exclude)
        except Exception as e:
            print(f"An error occurred: {e}")
            return None

    @staticmethod
    def AnalizeDocx(path, words, exclude):
        try:
            doc = Document(path)
            full_text = []

            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text = '\n'.join(full_text)
            return TextAnalizatorCore.AnalizeGeneric(text, words, exclude)
        except Exception as e:
            print(f"An error occurred: {e}")
            return None

    @staticmethod
    def GetWordsCount(words):
        text_without_punctuation = re.sub(r'[^\w\s]', '', words)
        return len(text_without_punctuation.split())

    @staticmethod
    def remove_specified_words_case_insensitive(input_string, words_to_remove):
        input_string_lower = input_string.lower()
        for word in words_to_remove:
            input_string_lower = input_string_lower.replace(word.lower(), '')
        return input_string_lower

    @staticmethod
    def count_specified_words(input_string, words_to_count):
        words = input_string.split()
        word_count = sum(word == specified_word for specified_word in words_to_count for word in words)
        return word_count
    
    @staticmethod
    def getDistribution(text):
        # Remove non-alphanumeric characters and split the text into words
        words = re.findall(r'\b\w+\b', text)

        # Count the occurrences of each word
        word_counts = Counter(words)

        # Find the most common words (at least 4 of them)
        most_common_words = word_counts.most_common(4)

        # Calculate the percentage of each most common word
        total_words = len(words)
        word_percentages = [(word, count, (count / total_words) * 100) for word, count in most_common_words]

        return word_percentages
    
    @staticmethod
    def count_words_with_marks(input_string):
    # Remove leading and trailing whitespaces
        cleaned_string = input_string.strip()

        # Remove punctuation marks
        cleaned_string = cleaned_string.translate(str.maketrans("", "", string.punctuation))

        # Split the string into words
        words = cleaned_string.split()

        # Count the number of words
        word_count = len(words)

        return word_count
    
    @staticmethod
    def countTextWater(text):
        language, confidence = langid.classify(text)
        if language == "ru":
            lan_file = TextAnalizatorCore.path_to_stopwords_file_ru
        elif language == "uk":
            lan_file = TextAnalizatorCore.path_to_stopwords_file_ua
        elif language == "en":
            lan_file = TextAnalizatorCore.path_to_stopwords_file_eng
        stopwords = []
        with open(lan_file, 'r', encoding='utf-8') as file:
            for line in file:
                stopwords.append(line.strip())
        return stopwords
    
    @staticmethod
    def AnalizeGeneric(text,words,exclude):
        try:
            stopWordsCount = TextAnalizatorCore.count_specified_words(text, exclude)
            specifiedCount = TextAnalizatorCore.count_specified_words(text,words)
            defaultText = text
            text = TextAnalizatorCore.remove_specified_words_case_insensitive(text, exclude)
            wordsCount = TextAnalizatorCore.count_words_with_marks(text)
            stopWords = TextAnalizatorCore.countTextWater(text)
            waterCount = TextAnalizatorCore.count_specified_words(text,stopWords)/wordsCount * 100
            symbolsCount = len(text)
            symbolsNoSpacesCount = len([char for char in text if not char.isspace()])
            lettersCount = sum(char.isalpha() for char in text)
            marksCount = sum(char in string.punctuation for char in text)
            dist = TextAnalizatorCore.getDistribution(text)
            text = TextAnalizatorCore.remove_specified_words_case_insensitive(text, stopWords)
            distNoStopWords = TextAnalizatorCore.getDistribution(text)
            return ResultModel( text=defaultText,
                                words=wordsCount,
                                specWords=specifiedCount, 
                                symbols=symbolsCount, 
                                symbolsNoSpaces=symbolsNoSpacesCount, 
                                letters=lettersCount, 
                                foreignWords = 0, 
                                waterPercentage=waterCount, 
                                marks=marksCount, 
                                stopWords=stopWordsCount, 
                                wordsDistribution=dist,
                                wordsDistributionNoStopWords=distNoStopWords) 
        except Exception as e:
            print(f"An error occurred: {e}")
            return None
        
        
            